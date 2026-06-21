from docx import Document

def export_word(data):
    doc = Document()

    doc.add_heading("Analisi Circolare", 0)

    doc.add_paragraph(data["sintesi"])

    doc.add_heading("Obblighi", 1)
    for o in data["obblighi"]:
        doc.add_paragraph(o)

    doc.save("output.docx")
